import os
from pathlib import Path
from typing import List, Dict, Tuple
from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd
import json
from datetime import datetime

class DocumentLoader:
    """Load documents from various formats (PDF, DOCX, CSV, XLSX)"""
    
    def __init__(self):
        self.documents = {}  # Store loaded documents
        self.metadata = {}   # Store metadata about documents
        
    def load_pdfs(self, pdf_dir: Path) -> Dict[str, str]:
        """Load all PDFs from a directory"""
        pdf_content = {}
        if not pdf_dir.exists():
            print(f"Warning: PDF directory not found: {pdf_dir}")
            return pdf_content
            
        for pdf_file in pdf_dir.glob("*.pdf"):
            try:
                text = ""
                with open(pdf_file, "rb") as f:
                    reader = PdfReader(f)
                    for page_num, page in enumerate(reader.pages):
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page.extract_text()
                
                pdf_content[pdf_file.name] = text
                self.metadata[pdf_file.name] = {
                    "type": "PDF",
                    "path": str(pdf_file),
                    "size": pdf_file.stat().st_size,
                    "loaded_at": datetime.now().isoformat()
                }
                print(f"Loaded PDF: {pdf_file.name}")
            except Exception as e:
                print(f"Error loading PDF {pdf_file.name}: {e}")
                
        return pdf_content
    
    def load_docx(self, docx_file: Path) -> str:
        """Load a single DOCX file"""
        try:
            doc = DocxDocument(docx_file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
            
            self.metadata[docx_file.name] = {
                "type": "DOCX",
                "path": str(docx_file),
                "size": docx_file.stat().st_size,
                "loaded_at": datetime.now().isoformat()
            }
            print(f"Loaded DOCX: {docx_file.name}")
            return text
        except Exception as e:
            print(f"Error loading DOCX {docx_file.name}: {e}")
            return ""
    
    def load_docx_files(self, docx_dir: Path) -> Dict[str, str]:
        """Load all DOCX files from a directory"""
        docx_content = {}
        if not docx_dir.exists():
            print(f"Warning: DOCX directory not found: {docx_dir}")
            return docx_content
            
        for docx_file in docx_dir.glob("*.docx"):
            text = self.load_docx(docx_file)
            if text:
                docx_content[docx_file.name] = text
                
        return docx_content
    
    def load_csv(self, csv_file: Path) -> Tuple[pd.DataFrame, Dict]:
        """Load a CSV file"""
        try:
            df = pd.read_csv(csv_file)
            self.metadata[csv_file.name] = {
                "type": "CSV",
                "path": str(csv_file),
                "rows": len(df),
                "columns": list(df.columns),
                "loaded_at": datetime.now().isoformat()
            }
            print(f"Loaded CSV: {csv_file.name} ({len(df)} rows, {len(df.columns)} columns)")
            return df, self.metadata[csv_file.name]
        except Exception as e:
            print(f"Error loading CSV {csv_file.name}: {e}")
            return pd.DataFrame(), {}
    
    def load_excel(self, excel_file: Path) -> Tuple[Dict[str, pd.DataFrame], Dict]:
        """Load an Excel file (all sheets)"""
        try:
            excel_data = pd.read_excel(excel_file, sheet_name=None)
            self.metadata[excel_file.name] = {
                "type": "XLSX",
                "path": str(excel_file),
                "sheets": list(excel_data.keys()),
                "loaded_at": datetime.now().isoformat()
            }
            print(f"Loaded XLSX: {excel_file.name} ({len(excel_data)} sheets)")
            return excel_data, self.metadata[excel_file.name]
        except Exception as e:
            print(f"Error loading XLSX {excel_file.name}: {e}")
            return {}, {}
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict]:
        """Split text into overlapping chunks for embedding"""
        chunks = []
        words = text.split()
        chunk_words = []
        
        for i, word in enumerate(words):
            chunk_words.append(word)
            if len(chunk_words) >= chunk_size // 5:  # Approximate words per chunk
                chunk_text = " ".join(chunk_words)
                chunks.append({
                    "text": chunk_text,
                    "word_count": len(chunk_words),
                    "position": i
                })
                # Keep last 'overlap' words
                chunk_words = chunk_words[-(overlap // 5):]
        
        # Add remaining text
        if chunk_words:
            chunks.append({
                "text": " ".join(chunk_words),
                "word_count": len(chunk_words),
                "position": len(words)
            })
        
        return chunks
    
    def get_metadata(self) -> Dict:
        """Return all loaded documents metadata"""
        return self.metadata
